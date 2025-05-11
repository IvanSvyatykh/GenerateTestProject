import os
from openpyxl import Workbook
from docx import Document
from fpdf import FPDF
from datetime import datetime
import json


async def generate_docx(test_json, theme, user_id, format_response):
    if isinstance(test_json, str):
        try:
            test_json = json.loads(test_json)
        except json.JSONDecodeError:
            raise ValueError("Некорректный JSON")

    if not isinstance(test_json, dict) or "questions" not in test_json:
        raise ValueError("test_json должен содержать ключ 'questions'")

    questions = test_json["questions"]
    document = Document()
    document.add_heading(f"Тема теста: {theme}", level=1)

    for idx, question in enumerate(questions, 1):
        q_text = f"{idx}. {question['question']}"
        document.add_paragraph(q_text)

        if (
            format_response == "✅ Варианты ответов"
            or format_response == "🔠 С пропусками"
        ):
            for index, answer in enumerate(question.get("answers", []), 1):
                document.add_paragraph(f"   {index}. {answer}")
        elif format_response == "🔀 Смешанный":
            if question.get("type") == "choice":
                for index, answer in enumerate(question.get("answers", []), 1):
                    document.add_paragraph(f"   {index}. {answer}")
            document.add_paragraph("")
        elif format_response == "📜 Открытые вопросы":
            document.add_paragraph("")

    document.add_page_break()
    document.add_heading("Ответы на тест:", level=1)

    for idx, question in enumerate(questions, 1):
        document.add_paragraph(f"{idx}. {question['correct_answer']}")

    file_path = f"tests/{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    os.makedirs("tests", exist_ok=True)
    document.save(file_path)

    return file_path


async def generate_pdf(test_json, theme, user_id, format_response):
    if isinstance(test_json, str):
        try:
            test_json = json.loads(test_json)
        except json.JSONDecodeError:
            raise ValueError("Некорректный JSON")

    if not isinstance(test_json, dict) or "questions" not in test_json:
        raise ValueError("test_json должен содержать ключ 'questions'")

    questions = test_json["questions"]
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("ArialUnicode", "", "ttf/Arial-Unicode-Regular.ttf", uni=True)
    pdf.set_font("ArialUnicode", size=12)

    # Установим ширину для переноса текста (например, 180)
    text_width = 180

    # Для заголовка
    pdf.multi_cell(text_width, 10, txt=f"Тема теста: {theme}", align="C")
    pdf.ln()  # Добавим пустую строку после заголовка

    for idx, question in enumerate(questions, 1):
        # Вопрос с автоматическим переносом
        pdf.multi_cell(text_width, 10, txt=f"{idx}. {question['question']}", align="J")
        pdf.ln(5)  # Отступ после вопроса

        if (
            format_response == "✅ Варианты ответов"
            or format_response == "🔠 С пропусками"
        ):
            for index, answer in enumerate(question.get("answers", []), 1):
                pdf.set_x(20)
                pdf.multi_cell(text_width - 20, 10, txt=f"{index}. {answer}")
        elif format_response == "🔀 Смешанный":
            if question.get("type") == "choice":
                for index, answer in enumerate(question.get("answers", []), 1):
                    pdf.set_x(20)
                    pdf.multi_cell(text_width - 20, 10, txt=f"{index}. {answer}")
            pdf.ln(5)
        elif format_response == "📜 Открытые вопросы":
            pdf.ln(5)

    pdf.add_page()
    pdf.multi_cell(text_width, 10, txt="Ответы на тест:", align="C")
    pdf.ln()

    for idx, question in enumerate(questions, 1):
        pdf.set_x(20)  # Отступ слева, чтобы не прилипало к краю
        pdf.multi_cell(text_width - 20, 10, txt=f"{idx}. {question['correct_answer']}")

    file_path = f"tests/{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    os.makedirs("tests", exist_ok=True)
    pdf.output(file_path)
    return file_path


async def generate_excel(test_json, theme, user_id, format_response):
    if isinstance(test_json, str):
        try:
            test_json = json.loads(test_json)
        except json.JSONDecodeError:
            raise ValueError("Некорректный JSON")

    if not isinstance(test_json, dict) or "questions" not in test_json:
        raise ValueError("test_json должен содержать ключ 'questions'")

    questions = test_json["questions"]
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Тест"
    sheet["A1"] = f"Тема теста: {theme}"
    row = 2

    for idx, question in enumerate(questions, 1):
        sheet.cell(row=row, column=1, value=f"{idx}. {question['question']}")
        row += 1

        if (
            format_response == "✅ Варианты ответов"
            or format_response == "🔠 С пропусками"
        ):
            for index, answer in enumerate(question.get("answers", []), 1):
                sheet.cell(row=row, column=1, value=f"   {index}. {answer}")
                row += 1
        elif format_response == "🔀 Смешанный":
            if question.get("type") == "choice":
                for index, answer in enumerate(question.get("answers", []), 1):
                    sheet.cell(row=row, column=1, value=f"   {index}. {answer}")
                    row += 1
            row += 1
        elif format_response == "📜 Открытые вопросы":
            row += 1

    row += 1
    sheet.cell(row=row, column=1, value="Ответы на тест:")
    row += 1
    for idx, question in enumerate(questions, 1):
        sheet.cell(row=row, column=1, value=f"{idx}. {question['correct_answer']}")
        row += 1

    file_path = f"tests/{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    os.makedirs("tests", exist_ok=True)
    workbook.save(file_path)
    return file_path
